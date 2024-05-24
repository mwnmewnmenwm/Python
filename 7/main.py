from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.label import Label
from kivy.uix.button import Button
from kivy.uix.spinner import Spinner
from kivy.uix.textinput import TextInput
from abc import ABC, abstractmethod
import docx

from recipes import wok, burger, pizza

class Recipe(ABC):
    def __init__(self):
        self._ingredients = {}

    @property
    def ingredients(self):
        return self._ingredients

    @abstractmethod
    def __str__(self):
        pass

    @abstractmethod
    def __repr__(self):
        pass

class Wok(Recipe):
    def __init__(self):
        super().__init__()
        self._ingredients = wok.ingredients

    def __str__(self):
        return "Wok Recipe"

    def __repr__(self):
        return f"Wok({self._ingredients})"

class Burger(Recipe):
    def __init__(self):
        super().__init__()
        self._ingredients = burger.ingredients

    def __str__(self):
        return "Burger Recipe"

    def __repr__(self):
        return f"Burger({self._ingredients})"

class Pizza(Recipe):
    def __init__(self):
        super().__init__()
        self._ingredients = pizza.ingredients

    def __str__(self):
        return "Pizza Recipe"

    def __repr__(self):
        return f"Pizza({self._ingredients})"

class RecipeCalculator(BoxLayout):
    def __init__(self, **kwargs):
        super(RecipeCalculator, self).__init__(**kwargs)

        self.orientation = 'vertical'

        self.recipe_label = Label(text="Выберите рецепт:")
        self.add_widget(self.recipe_label)

        self.recipe_spinner = Spinner(
            text='Pizza',
            values=('Wok', 'Burger', 'Pizza')
        )
        self.recipe_spinner.bind(text=self.update_ingredients_on_change)
        self.add_widget(self.recipe_spinner)

        self.ingredients_frame = BoxLayout(orientation='vertical')
        self.add_widget(self.ingredients_frame)

        self.calculate_button = Button(text="Рассчитать")
        self.calculate_button.bind(on_press=self.calculate)
        self.add_widget(self.calculate_button)

        self.result_label = Label(text="")
        self.add_widget(self.result_label)

        self.recipe_instances = {
            "Wok": Wok(),
            "Burger": Burger(),
            "Pizza": Pizza()
        }

        self.update_ingredients(self.recipe_instances["Pizza"].ingredients)

    def update_ingredients_on_change(self, instance, value):
        selected_recipe = self.recipe_instances[value].ingredients
        self.update_ingredients(selected_recipe)

    def update_ingredients(self, selected_recipe):
        self.ingredients_frame.clear_widgets()
        self.ingredient_entries = {}

        for ingredient in selected_recipe.keys():
            label = Label(text=ingredient)
            self.ingredients_frame.add_widget(label)
            entry = TextInput()
            self.ingredients_frame.add_widget(entry)
            self.ingredient_entries[ingredient] = entry

    def calculate(self, instance):
        selected_recipe_name = self.recipe_spinner.text
        selected_recipe_ingredients = self.recipe_instances[selected_recipe_name].ingredients

        total_calories = 0
        total_cost = 0

        for ingredient, info in selected_recipe_ingredients.items():
            try:
                quantity = float(self.ingredient_entries[ingredient].text)
                total_calories += info['calories'] * (quantity / 100)
                total_cost += info['price'] * (quantity / 100)
            except ValueError:
                self.result_label.text = "Пожалуйста, введите корректное количество для всех ингредиентов."
                return

        result_text = f"Выбранный рецепт: {selected_recipe_name}\n"
        result_text += f"Энергетическая ценность: {total_calories} кал\n"
        result_text += f"Стоимость рецепта: {total_cost} руб"
        self.result_label.text = result_text

        self.save_report(selected_recipe_name, total_calories, total_cost)

    def save_report(self, selected_recipe, calories, cost):
        doc = docx.Document()
        doc.add_heading('Отчет о расчете рецепта', level=1)
        doc.add_paragraph(f"Выбранный рецепт: {selected_recipe}")
        doc.add_paragraph(f"Энергетическая ценность: {calories} кал")
        doc.add_paragraph(f"Стоимость рецепта: {cost} руб")
        doc.save('отчет_рецепта.docx')

class RecipeCalculatorApp(App):
    def build(self):
        return RecipeCalculator()

if __name__ == '__main__':
    RecipeCalculatorApp().run()
