from docx import Document

def create_docx_report(recipe_name, recipe, total_calories, total_cost):
    doc = Document() # создаем документ
    doc.add_heading(f'{recipe_name}', 0) 
    for ingredient, quantity in recipe.items():
        doc.add_heading(f'{ingredient} - {quantity}г', level=1)
    doc.add_paragraph(f'Энергетическая ценность: {total_calories} ккал')
    doc.add_paragraph(f'Стоимость: {total_cost} руб.')
    doc.save('ворд.docx')
