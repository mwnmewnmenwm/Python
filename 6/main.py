import tkinter as tk
from tkinter import messagebox
from recipes import wok, burger, pizza
import docx

def calculate():
    selected_recipe = None
    if recipe_var.get() == "Wok":
        selected_recipe = wok.ingredients
    elif recipe_var.get() == "Burger":
        selected_recipe = burger.ingredients
    elif recipe_var.get() == "Pizza":
        selected_recipe = pizza.ingredients
    
    total_calories = 0
    total_cost = 0
    for ingredient, info in selected_recipe.items():
        quantity = float(ingredient_entries[ingredient].get())
        total_calories += info['calories'] * (quantity / 100)
        total_cost += info['price'] * (quantity / 100)
    
    result_text = f"Выбранный рецепт: {recipe_var.get()}\n"
    result_text += f"Энергетическая ценность: {total_calories} кал\n"
    result_text += f"Стоимость рецепта: {total_cost} руб"
    result_label.config(text=result_text)

    save_report(recipe_var.get(), total_calories, total_cost)

def save_report(selected_recipe, calories, cost):
    doc = docx.Document()
    doc.add_heading('Отчет о расчете рецепта', level=1)
    doc.add_paragraph(f"Выбранный рецепт: {selected_recipe}")
    doc.add_paragraph(f"Энергетическая ценность: {calories} кал")
    doc.add_paragraph(f"Стоимость рецепта: {cost} руб")
    doc.save('ворд.docx')

def update_ingredients(selected_recipe):
    for widget in ingredients_frame.winfo_children():
        widget.destroy()
    
    global ingredient_entries
    ingredient_entries = {}
    for ingredient in selected_recipe.keys():
        label = tk.Label(ingredients_frame, text=ingredient)
        label.pack(side=tk.LEFT)
        entry = tk.Entry(ingredients_frame)
        entry.pack(side=tk.LEFT)
        ingredient_entries[ingredient] = entry

root = tk.Tk()
root.title("Recipe Calculator")

recipe_label = tk.Label(root, text="Выберите рецепт:")
recipe_label.pack()

recipe_var = tk.StringVar(root)
recipe_var.set("Wok")

recipe_options = ["Wok", "Burger", "Pizza"]

def update_ingredients_on_change(*args):
    selected_recipe = None
    if recipe_var.get() == "Wok":
        selected_recipe = wok.ingredients
    elif recipe_var.get() == "Burger":
        selected_recipe = burger.ingredients
    elif recipe_var.get() == "Pizza":
        selected_recipe = pizza.ingredients
    
    update_ingredients(selected_recipe)

recipe_var.trace("w", update_ingredients_on_change)

recipe_menu = tk.OptionMenu(root, recipe_var, *recipe_options)
recipe_menu.pack()

ingredients_frame = tk.Frame(root)
ingredients_frame.pack()

ingredient_entries = {}
update_ingredients(wok.ingredients)

calculate_button = tk.Button(root, text="Рассчитать", command=calculate)
calculate_button.pack()

result_label = tk.Label(root, text="")
result_label.pack()

root.mainloop()
